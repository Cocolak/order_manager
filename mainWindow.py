from PyQt5 import QtCore, QtGui, QtWidgets, Qt
from PyQt5.QtWidgets import QMessageBox, QDialog
import pandas as pd
import res


class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        self.MainWindow = MainWindow
        self.currentOrderID = None

        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 500)

        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")

        # Panel
        self.panelWidget = QtWidgets.QWidget(self.centralwidget)
        self.panelWidget.setMaximumSize(QtCore.QSize(150, 16777215))
        self.panelWidget.setObjectName("panelWidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.panelWidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.panel_searchLineEdit = QtWidgets.QLineEdit(self.panelWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.panel_searchLineEdit.sizePolicy().hasHeightForWidth())
        self.panel_searchLineEdit.setSizePolicy(sizePolicy)
        self.panel_searchLineEdit.setMaximumSize(QtCore.QSize(16777215, 30))
        self.panel_searchLineEdit.textChanged.connect(self.loadListWidget)
        self.panel_searchLineEdit.setObjectName("panel_searchLineEdit")
        self.verticalLayout.addWidget(self.panel_searchLineEdit)
        self.panel_listWidget = QtWidgets.QListWidget(self.panelWidget)
        self.panel_listWidget.itemClicked.connect(self.panel_openInfoWindow)
        self.panel_listWidget.setObjectName("panel_listWidget")
        self.verticalLayout.addWidget(self.panel_listWidget)
        self.panel_addButton = QtWidgets.QPushButton(self.panelWidget)
        self.panel_addButton.setObjectName("panel_addButton")
        self.panel_addButton.clicked.connect(self.panel_addButtonClicked)
        self.verticalLayout.addWidget(self.panel_addButton)
        self.panel_excelPreview = QtWidgets.QPushButton(self.panelWidget)
        self.panel_excelPreview.setObjectName("panel_excelPreview")
        self.panel_excelPreview.clicked.connect(self.panel_excelPreviewClicked)
        self.verticalLayout.addWidget(self.panel_excelPreview)
        self.gridLayout.addWidget(self.panelWidget, 0, 0, 2, 1)

        # Top Widget
        self.topWidget = QtWidgets.QWidget(self.centralwidget)
        self.topWidget.setObjectName("topWidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.topWidget)
        self.horizontalLayout.setObjectName("horizontalLayout")
        spacerItem = QtWidgets.QSpacerItem(153, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout.addItem(spacerItem)
        self.top_actualLabel = QtWidgets.QLabel(self.topWidget)
        font = QtGui.QFont()
        font.setPointSize(16)
        self.top_actualLabel.setFont(font)
        self.top_actualLabel.setObjectName("top_actualLabel")
        self.horizontalLayout.addWidget(self.top_actualLabel)
        spacerItem1 = QtWidgets.QSpacerItem(153, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout.addItem(spacerItem1)
        self.top_logoLabel = QtWidgets.QLabel(self.topWidget)
        self.top_logoLabel.setMinimumSize(QtCore.QSize(120, 40))
        self.top_logoLabel.setMaximumSize(QtCore.QSize(100, 30))
        self.top_logoLabel.setText("")
        self.top_logoLabel.setPixmap(QtGui.QPixmap(":/images/images/mfcomp.png"))
        self.top_logoLabel.setScaledContents(True)
        self.top_logoLabel.setObjectName("top_logoLabel")
        self.horizontalLayout.addWidget(self.top_logoLabel)
        self.gridLayout.addWidget(self.topWidget, 0, 1, 1, 1)

        # Pages
        self.stackedWidget = QtWidgets.QStackedWidget(self.centralwidget)
        self.stackedWidget.setObjectName("stackedWidget")

            # homePage
        self.homePage = QtWidgets.QWidget()
        self.homePage.setObjectName("homePage")
        self.stackedWidget.addWidget(self.homePage)

            # infoPage
        self.infoPage = QtWidgets.QWidget()
        self.infoPage.setObjectName("infoPage")
        self.verticalLayout_7 = QtWidgets.QVBoxLayout(self.infoPage)
        self.verticalLayout_7.setObjectName("verticalLayout_7")
        self.infoContentWidget = QtWidgets.QWidget(self.infoPage)
        self.infoContentWidget.setObjectName("infoContentWidget")
        self.gridLayout_3 = QtWidgets.QGridLayout(self.infoContentWidget)
        self.gridLayout_3.setHorizontalSpacing(20)
        self.gridLayout_3.setVerticalSpacing(10)
        self.gridLayout_3.setObjectName("gridLayout_3")
        self.info_dateLayout = QtWidgets.QHBoxLayout()
        self.info_dateLayout.setSpacing(8)
        self.info_dateLayout.setObjectName("info_dateLayout")
        self.info_dateLabel = QtWidgets.QLabel(self.infoContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_dateLabel.sizePolicy().hasHeightForWidth())
        self.info_dateLabel.setSizePolicy(sizePolicy)
        self.info_dateLabel.setMinimumSize(QtCore.QSize(0, 30))
        self.info_dateLabel.setObjectName("info_dateLabel")
        self.info_dateLayout.addWidget(self.info_dateLabel, 0, QtCore.Qt.AlignVCenter)
        self.info_dateContentLabel = QtWidgets.QLabel(self.infoContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_dateContentLabel.sizePolicy().hasHeightForWidth())
        self.info_dateContentLabel.setSizePolicy(sizePolicy)
        self.info_dateContentLabel.setObjectName("info_dateContentLabel")
        self.info_dateLayout.addWidget(self.info_dateContentLabel)
        self.gridLayout_3.addLayout(self.info_dateLayout, 0, 1, 1, 1)
        self.info_usLayout = QtWidgets.QVBoxLayout()
        self.info_usLayout.setSpacing(8)
        self.info_usLayout.setObjectName("info_usLayout")
        self.info_usLabel = QtWidgets.QLabel(self.infoContentWidget)
        self.info_usLabel.setObjectName("info_usLabel")
        self.info_usLayout.addWidget(self.info_usLabel)
        self.info_usTextEdit = QtWidgets.QTextEdit(self.infoContentWidget)
        self.info_usTextEdit.setMinimumSize(QtCore.QSize(0, 150))
        self.info_usTextEdit.setReadOnly(True)
        self.info_usTextEdit.setObjectName("info_usTextEdit")
        self.info_usLayout.addWidget(self.info_usTextEdit)
        self.gridLayout_3.addLayout(self.info_usLayout, 2, 1, 1, 1)
        self.info_modelLayout = QtWidgets.QHBoxLayout()
        self.info_modelLayout.setSpacing(8)
        self.info_modelLayout.setObjectName("info_modelLayout")
        self.info_modelLabel = QtWidgets.QLabel(self.infoContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_modelLabel.sizePolicy().hasHeightForWidth())
        self.info_modelLabel.setSizePolicy(sizePolicy)
        self.info_modelLabel.setMinimumSize(QtCore.QSize(0, 30))
        self.info_modelLabel.setObjectName("info_modelLabel")
        self.info_modelLayout.addWidget(self.info_modelLabel, 0, QtCore.Qt.AlignVCenter)
        self.info_modelContentLabel = QtWidgets.QLabel(self.infoContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_modelContentLabel.sizePolicy().hasHeightForWidth())
        self.info_modelContentLabel.setSizePolicy(sizePolicy)
        self.info_modelContentLabel.setObjectName("info_modelContentLabel")
        self.info_modelLayout.addWidget(self.info_modelContentLabel)
        self.gridLayout_3.addLayout(self.info_modelLayout, 1, 0, 1, 1)
        self.info_descLayout = QtWidgets.QVBoxLayout()
        self.info_descLayout.setSpacing(8)
        self.info_descLayout.setObjectName("info_descLayout")
        self.info_descLabel = QtWidgets.QLabel(self.infoContentWidget)
        self.info_descLabel.setObjectName("info_descLabel")
        self.info_descLayout.addWidget(self.info_descLabel)
        self.info_descTextEdit = QtWidgets.QTextEdit(self.infoContentWidget)
        self.info_descTextEdit.setMinimumSize(QtCore.QSize(0, 150))
        self.info_descTextEdit.setReadOnly(True)
        self.info_descTextEdit.setObjectName("info_descTextEdit")
        self.info_descLayout.addWidget(self.info_descTextEdit)
        self.gridLayout_3.addLayout(self.info_descLayout, 2, 0, 1, 1)
        self.info_nrLayout = QtWidgets.QHBoxLayout()
        self.info_nrLayout.setSpacing(8)
        self.info_nrLayout.setObjectName("info_nrLayout")
        self.info_nrLabel = QtWidgets.QLabel(self.infoContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_nrLabel.sizePolicy().hasHeightForWidth())
        self.info_nrLabel.setSizePolicy(sizePolicy)
        self.info_nrLabel.setMinimumSize(QtCore.QSize(0, 30))
        self.info_nrLabel.setObjectName("info_nrLabel")
        self.info_nrLayout.addWidget(self.info_nrLabel, 0, QtCore.Qt.AlignVCenter)
        self.info_nrContentLabel = QtWidgets.QLabel(self.infoContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_nrContentLabel.sizePolicy().hasHeightForWidth())
        self.info_nrContentLabel.setSizePolicy(sizePolicy)
        self.info_nrContentLabel.setObjectName("info_nrContentLabel")
        self.info_nrLayout.addWidget(self.info_nrContentLabel)
        self.gridLayout_3.addLayout(self.info_nrLayout, 0, 0, 1, 1)
        self.info_nrtelLayout = QtWidgets.QHBoxLayout()
        self.info_nrtelLayout.setSpacing(8)
        self.info_nrtelLayout.setObjectName("info_nrtelLayout")
        self.info_nrtelLabel = QtWidgets.QLabel(self.infoContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_nrtelLabel.sizePolicy().hasHeightForWidth())
        self.info_nrtelLabel.setSizePolicy(sizePolicy)
        self.info_nrtelLabel.setMinimumSize(QtCore.QSize(0, 30))
        self.info_nrtelLabel.setObjectName("info_nrtelLabel")
        self.info_nrtelLayout.addWidget(self.info_nrtelLabel, 0, QtCore.Qt.AlignVCenter)
        self.info_nrtelContentLabel = QtWidgets.QLabel(self.infoContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_nrtelContentLabel.sizePolicy().hasHeightForWidth())
        self.info_nrtelContentLabel.setSizePolicy(sizePolicy)
        self.info_nrtelContentLabel.setObjectName("info_nrtelContentLabel")
        self.info_nrtelLayout.addWidget(self.info_nrtelContentLabel)
        self.gridLayout_3.addLayout(self.info_nrtelLayout, 1, 1, 1, 1)
        self.verticalLayout_7.addWidget(self.infoContentWidget)
        self.infoButtonsWidget = QtWidgets.QWidget(self.infoPage)
        self.infoButtonsWidget.setObjectName("infoButtonsWidget")
        self.horizontalLayout_8 = QtWidgets.QHBoxLayout(self.infoButtonsWidget)
        self.horizontalLayout_8.setObjectName("horizontalLayout_8")
        self.info_isReadyCheckBox = QtWidgets.QCheckBox(self.infoButtonsWidget)
        self.info_isReadyCheckBox.setObjectName("info_isReadyCheckBox")
        self.horizontalLayout_8.addWidget(self.info_isReadyCheckBox)
        spacerItem2 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_8.addItem(spacerItem2)
        self.info_printButton = QtWidgets.QPushButton(self.infoButtonsWidget)
        self.info_printButton.setMinimumSize(QtCore.QSize(100, 30))
        self.info_printButton.clicked.connect(self.info_printButtonClicked)
        self.info_printButton.setObjectName("info_printButton")
        self.horizontalLayout_8.addWidget(self.info_printButton)
        self.info_editButton = QtWidgets.QPushButton(self.infoButtonsWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_editButton.sizePolicy().hasHeightForWidth())
        self.info_editButton.setSizePolicy(sizePolicy)
        self.info_editButton.setMinimumSize(QtCore.QSize(100, 30))
        self.info_editButton.clicked.connect(self.editButtonClicked)
        self.info_editButton.setObjectName("info_editButton")
        self.horizontalLayout_8.addWidget(self.info_editButton)
        self.info_removeButton = QtWidgets.QPushButton(self.infoButtonsWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.info_removeButton.sizePolicy().hasHeightForWidth())
        self.info_removeButton.setSizePolicy(sizePolicy)
        self.info_removeButton.setMinimumSize(QtCore.QSize(100, 30))
        self.info_removeButton.clicked.connect(self.info_removeButtonClicked)
        self.info_removeButton.setObjectName("info_removeButton")
        self.horizontalLayout_8.addWidget(self.info_removeButton)
        spacerItem3 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_8.addItem(spacerItem3)
        self.verticalLayout_7.addWidget(self.infoButtonsWidget)
        self.stackedWidget.addWidget(self.infoPage)

            # editPage
        self.editPage = QtWidgets.QWidget()
        self.editPage.setObjectName("editPage")
        self.verticalLayout_10 = QtWidgets.QVBoxLayout(self.editPage)
        self.verticalLayout_10.setObjectName("verticalLayout_10")
        self.editContentWidget = QtWidgets.QWidget(self.editPage)
        self.editContentWidget.setObjectName("editContentWidget")
        self.gridLayout_4 = QtWidgets.QGridLayout(self.editContentWidget)
        self.gridLayout_4.setHorizontalSpacing(20)
        self.gridLayout_4.setVerticalSpacing(16)
        self.gridLayout_4.setObjectName("gridLayout_4")
        self.edit_dateLayout = QtWidgets.QHBoxLayout()
        self.edit_dateLayout.setSpacing(8)
        self.edit_dateLayout.setObjectName("edit_dateLayout")
        self.edit_dateLabel = QtWidgets.QLabel(self.editContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.edit_dateLabel.sizePolicy().hasHeightForWidth())
        self.edit_dateLabel.setSizePolicy(sizePolicy)
        self.edit_dateLabel.setObjectName("edit_dateLabel")
        self.edit_dateLayout.addWidget(self.edit_dateLabel, 0, QtCore.Qt.AlignVCenter)
        self.edit_dateLineEdit = QtWidgets.QLineEdit(self.editContentWidget)
        self.edit_dateLineEdit.setReadOnly(True)
        self.edit_dateLineEdit.setObjectName("edit_dateLineEdit")
        self.edit_dateLayout.addWidget(self.edit_dateLineEdit)
        self.gridLayout_4.addLayout(self.edit_dateLayout, 0, 1, 1, 1)
        self.edit_usLayout = QtWidgets.QVBoxLayout()
        self.edit_usLayout.setSpacing(8)
        self.edit_usLayout.setObjectName("edit_usLayout")
        self.edit_usLabel = QtWidgets.QLabel(self.editContentWidget)
        self.edit_usLabel.setObjectName("edit_usLabel")
        self.edit_usLayout.addWidget(self.edit_usLabel)
        self.edit_usTextEdit = QtWidgets.QTextEdit(self.editContentWidget)
        self.edit_usTextEdit.setMinimumSize(QtCore.QSize(0, 150))
        self.edit_usTextEdit.setObjectName("edit_usTextEdit")
        self.edit_usLayout.addWidget(self.edit_usTextEdit)
        self.gridLayout_4.addLayout(self.edit_usLayout, 2, 1, 1, 1)
        self.edit_modelLayout = QtWidgets.QHBoxLayout()
        self.edit_modelLayout.setSpacing(8)
        self.edit_modelLayout.setObjectName("edit_modelLayout")
        self.edit_modelLabel = QtWidgets.QLabel(self.editContentWidget)
        self.edit_modelLabel.setObjectName("edit_modelLabel")
        self.edit_modelLayout.addWidget(self.edit_modelLabel, 0, QtCore.Qt.AlignVCenter)
        self.edit_modelLineEdit = QtWidgets.QLineEdit(self.editContentWidget)
        self.edit_modelLineEdit.setObjectName("edit_modelLineEdit")
        self.edit_modelLayout.addWidget(self.edit_modelLineEdit)
        self.gridLayout_4.addLayout(self.edit_modelLayout, 1, 0, 1, 1)
        self.edit_descLayout = QtWidgets.QVBoxLayout()
        self.edit_descLayout.setSpacing(8)
        self.edit_descLayout.setObjectName("edit_descLayout")
        self.edit_descLabel = QtWidgets.QLabel(self.editContentWidget)
        self.edit_descLabel.setObjectName("edit_descLabel")
        self.edit_descLayout.addWidget(self.edit_descLabel)
        self.edit_descTextEdit = QtWidgets.QTextEdit(self.editContentWidget)
        self.edit_descTextEdit.setMinimumSize(QtCore.QSize(0, 150))
        self.edit_descTextEdit.setObjectName("edit_descTextEdit")
        self.edit_descLayout.addWidget(self.edit_descTextEdit)
        self.gridLayout_4.addLayout(self.edit_descLayout, 2, 0, 1, 1)
        self.edit_nrLayout = QtWidgets.QHBoxLayout()
        self.edit_nrLayout.setSpacing(8)
        self.edit_nrLayout.setObjectName("edit_nrLayout")
        self.edit_nrLabel = QtWidgets.QLabel(self.editContentWidget)
        self.edit_nrLabel.setObjectName("edit_nrLabel")
        self.edit_nrLayout.addWidget(self.edit_nrLabel, 0, QtCore.Qt.AlignVCenter)
        self.edit_nrLineEdit = QtWidgets.QLineEdit(self.editContentWidget)
        self.edit_nrLineEdit.setReadOnly(True)
        self.edit_nrLineEdit.setObjectName("edit_nrLineEdit")
        self.edit_nrLayout.addWidget(self.edit_nrLineEdit)
        self.gridLayout_4.addLayout(self.edit_nrLayout, 0, 0, 1, 1)
        self.edit_nrtelLayout = QtWidgets.QHBoxLayout()
        self.edit_nrtelLayout.setSpacing(8)
        self.edit_nrtelLayout.setObjectName("edit_nrtelLayout")
        self.edit_nrtelLabel = QtWidgets.QLabel(self.editContentWidget)
        self.edit_nrtelLabel.setObjectName("edit_nrtelLabel")
        self.edit_nrtelLayout.addWidget(self.edit_nrtelLabel, 0, QtCore.Qt.AlignVCenter)
        self.edit_nrtelLineEdit = QtWidgets.QLineEdit(self.editContentWidget)
        self.edit_nrtelLineEdit.setObjectName("edit_nrtelLineEdit")
        self.edit_nrtelLayout.addWidget(self.edit_nrtelLineEdit)
        self.gridLayout_4.addLayout(self.edit_nrtelLayout, 1, 1, 1, 1)
        self.verticalLayout_10.addWidget(self.editContentWidget)
        self.editButtonsWidget = QtWidgets.QWidget(self.editPage)
        self.editButtonsWidget.setObjectName("editbuttonsWidget")
        self.horizontalLayout_15 = QtWidgets.QHBoxLayout(self.editButtonsWidget)
        self.horizontalLayout_15.setObjectName("horizontalLayout_15")
        self.edit_isReadyCheckBox = QtWidgets.QCheckBox(self.editButtonsWidget)
        self.edit_isReadyCheckBox.setObjectName("edit_isReadyCheckBox")
        self.horizontalLayout_15.addWidget(self.edit_isReadyCheckBox)
        spacerItem4 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_15.addItem(spacerItem4)
        self.edit_saveButton = QtWidgets.QPushButton(self.editButtonsWidget)
        self.edit_saveButton.setMinimumSize(QtCore.QSize(100, 30))
        self.edit_saveButton.clicked.connect(self.edit_saveButtonClicked)
        self.edit_saveButton.setObjectName("edit_saveButton")
        self.horizontalLayout_15.addWidget(self.edit_saveButton)
        self.edit_cancelButton = QtWidgets.QPushButton(self.editButtonsWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.edit_cancelButton.sizePolicy().hasHeightForWidth())
        self.edit_cancelButton.setSizePolicy(sizePolicy)
        self.edit_cancelButton.setMinimumSize(QtCore.QSize(100, 30))
        self.edit_cancelButton.clicked.connect(self.edit_cancelButtonClicked)
        self.edit_cancelButton.setObjectName("edit_cancelButton")
        self.horizontalLayout_15.addWidget(self.edit_cancelButton)
        spacerItem5 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_15.addItem(spacerItem5)
        self.verticalLayout_10.addWidget(self.editButtonsWidget)
        self.stackedWidget.addWidget(self.editPage)

            # addPage
        self.addPage = QtWidgets.QWidget()
        self.addPage.setObjectName("addPage")
        self.verticalLayout_4 = QtWidgets.QVBoxLayout(self.addPage)
        self.verticalLayout_4.setObjectName("verticalLayout_4")
        self.addContentWidget = QtWidgets.QWidget(self.addPage)
        self.addContentWidget.setObjectName("addContentWidget")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.addContentWidget)
        self.gridLayout_2.setHorizontalSpacing(20)
        self.gridLayout_2.setVerticalSpacing(16)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.add_nrtelLayout = QtWidgets.QHBoxLayout()
        self.add_nrtelLayout.setSpacing(8)
        self.add_nrtelLayout.setObjectName("add_nrtelLayout")
        self.add_nrtelLabel = QtWidgets.QLabel(self.addContentWidget)
        self.add_nrtelLabel.setObjectName("add_nrtelLabel")
        self.add_nrtelLayout.addWidget(self.add_nrtelLabel, 0, QtCore.Qt.AlignVCenter)
        self.add_nrtelLineEdit = QtWidgets.QLineEdit(self.addContentWidget)
        self.add_nrtelLineEdit.setObjectName("add_nrtelLineEdit")
        self.add_nrtelLayout.addWidget(self.add_nrtelLineEdit)
        self.gridLayout_2.addLayout(self.add_nrtelLayout, 0, 0, 1, 1)
        self.add_modelLayout = QtWidgets.QHBoxLayout()
        self.add_modelLayout.setSpacing(8)
        self.add_modelLayout.setObjectName("add_modelLayout")
        self.add_modelLabel = QtWidgets.QLabel(self.addContentWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.add_modelLabel.sizePolicy().hasHeightForWidth())
        self.add_modelLabel.setSizePolicy(sizePolicy)
        self.add_modelLabel.setObjectName("add_modelLabel")
        self.add_modelLayout.addWidget(self.add_modelLabel, 0, QtCore.Qt.AlignVCenter)
        self.add_modelLineEdit = QtWidgets.QLineEdit(self.addContentWidget)
        self.add_modelLineEdit.setObjectName("add_modelLineEdit")
        self.add_modelLayout.addWidget(self.add_modelLineEdit)
        self.gridLayout_2.addLayout(self.add_modelLayout, 0, 1, 1, 1)
        self.add_descLayout = QtWidgets.QVBoxLayout()
        self.add_descLayout.setSpacing(8)
        self.add_descLayout.setObjectName("add_descLayout")
        self.add_descLabel = QtWidgets.QLabel(self.addContentWidget)
        self.add_descLabel.setObjectName("add_descLabel")
        self.add_descLayout.addWidget(self.add_descLabel)
        self.add_descTextEdit = QtWidgets.QTextEdit(self.addContentWidget)
        self.add_descTextEdit.setMinimumSize(QtCore.QSize(0, 200))
        self.add_descTextEdit.setObjectName("add_descTextEdit")
        self.add_descLayout.addWidget(self.add_descTextEdit)
        self.gridLayout_2.addLayout(self.add_descLayout, 1, 0, 1, 1)
        self.add_usLayout = QtWidgets.QVBoxLayout()
        self.add_usLayout.setSpacing(8)
        self.add_usLayout.setObjectName("add_usLayout")
        self.add_usLabel = QtWidgets.QLabel(self.addContentWidget)
        self.add_usLabel.setObjectName("add_usLabel")
        self.add_usLayout.addWidget(self.add_usLabel)
        self.add_usTextEdit = QtWidgets.QTextEdit(self.addContentWidget)
        self.add_usTextEdit.setMinimumSize(QtCore.QSize(0, 200))
        self.add_usTextEdit.setObjectName("add_usTextEdit")
        self.add_usLayout.addWidget(self.add_usTextEdit)
        self.gridLayout_2.addLayout(self.add_usLayout, 1, 1, 1, 1)
        self.verticalLayout_4.addWidget(self.addContentWidget)
        self.addButtonsWidget = QtWidgets.QWidget(self.addPage)
        self.addButtonsWidget.setObjectName("addButtonsWidget")
        self.horizontalLayout_3 = QtWidgets.QHBoxLayout(self.addButtonsWidget)
        self.horizontalLayout_3.setObjectName("horizontalLayout_3")
        spacerItem6 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_3.addItem(spacerItem6)
        self.add_confirmButton = QtWidgets.QPushButton(self.addButtonsWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.add_confirmButton.sizePolicy().hasHeightForWidth())
        self.add_confirmButton.setSizePolicy(sizePolicy)
        self.add_confirmButton.setMinimumSize(QtCore.QSize(100, 30))
        self.add_confirmButton.setObjectName("add_confirmButton")
        self.add_confirmButton.clicked.connect(self.add_confirmClicked)
        self.horizontalLayout_3.addWidget(self.add_confirmButton)
        self.add_cancelButton = QtWidgets.QPushButton(self.addButtonsWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.add_cancelButton.sizePolicy().hasHeightForWidth())
        self.add_cancelButton.setSizePolicy(sizePolicy)
        self.add_cancelButton.setMinimumSize(QtCore.QSize(100, 30))
        self.add_cancelButton.clicked.connect(self.add_cancelClicked)
        self.add_cancelButton.setObjectName("add_cancelButton")
        self.horizontalLayout_3.addWidget(self.add_cancelButton)
        spacerItem7 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout_3.addItem(spacerItem7)
        self.verticalLayout_4.addWidget(self.addButtonsWidget)

        # ID widgets
        # 0 - homePage
        # 1 - infoPage
        # 2 - editPage
        # 3 - addPage
        self.stackedWidget.addWidget(self.addPage)
        self.gridLayout.addWidget(self.stackedWidget, 1, 1, 1, 1)
        MainWindow.setCentralWidget(self.centralwidget)

        self.loadListWidget()

        self.retranslateUi(MainWindow)
        self.stackedWidget.setCurrentIndex(0)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.panel_searchLineEdit.setPlaceholderText(_translate("MainWindow", "Szukaj..."))
        self.panel_addButton.setText(_translate("MainWindow", "DODAJ ZLECENIE"))
        self.panel_excelPreview.setText(_translate("MainWindow", "PODGLĄD EXCELA"))
        self.top_actualLabel.setText(_translate("MainWindow", "Strona główna"))
        self.info_dateLabel.setText(_translate("MainWindow", "Data przyjęcia:"))
        self.info_dateContentLabel.setText(_translate("MainWindow", "-"))
        self.info_usLabel.setText(_translate("MainWindow", "Dla nas:"))
        self.info_modelLabel.setText(_translate("MainWindow", "Model:"))
        self.info_modelContentLabel.setText(_translate("MainWindow", "-"))
        self.info_descLabel.setText(_translate("MainWindow", "Opis zlecenia:"))
        self.info_nrLabel.setText(_translate("MainWindow", "Nr. Zlecenia:"))
        self.info_nrContentLabel.setText(_translate("MainWindow", "-"))
        self.info_nrtelLabel.setText(_translate("MainWindow", "Nr. Telefonu:"))
        self.info_nrtelContentLabel.setText(_translate("MainWindow", "-"))
        self.info_isReadyCheckBox.setText(_translate("MainWindow", "Gotowe"))
        self.info_printButton.setText(_translate("MainWindow", "DRUKUJ"))
        self.info_editButton.setText(_translate("MainWindow", "EDYTUJ"))
        self.info_removeButton.setText(_translate("MainWindow", "USUŃ"))
        self.edit_dateLabel.setText(_translate("MainWindow", "Data przyjęcia:"))
        self.edit_usLabel.setText(_translate("MainWindow", "Dla nas:"))
        self.edit_modelLabel.setText(_translate("MainWindow", "Model:"))
        self.edit_descLabel.setText(_translate("MainWindow", "Opis zlecenia:"))
        self.edit_nrLabel.setText(_translate("MainWindow", "Nr. Zlecenia:"))
        self.edit_nrtelLabel.setText(_translate("MainWindow", "Nr. Telefonu:"))
        self.edit_isReadyCheckBox.setText(_translate("MainWindow", "Gotowe"))
        self.edit_saveButton.setText(_translate("MainWindow", "ZAPISZ"))
        self.edit_cancelButton.setText(_translate("MainWindow", "ANULUJ"))
        self.add_nrtelLabel.setText(_translate("MainWindow", "Nr. Telefonu"))
        self.add_modelLabel.setText(_translate("MainWindow", "Model:"))
        self.add_descLabel.setText(_translate("MainWindow", "Opis zlecenia:"))
        self.add_usLabel.setText(_translate("MainWindow", "Dla nas:"))
        self.add_confirmButton.setText(_translate("MainWindow", "POTWIERDŹ"))
        self.add_cancelButton.setText(_translate("MainWindow", "ANULUJ"))

    def loadListWidget(self):
        self.panel_listWidget.clear()

        #TODO zwiększ wielkość elementów listy
        #TODO: gdy na początku "+" szukaj po numerze tel.

        # Load data from file
        df = pd.read_excel("_dane/dane.xlsx")
        sorted_df = df.sort_values(["Zlecenie"], ascending=[False])
        ordersIDs = sorted_df["Zlecenie"].tolist() # Orders IDs descending sorted
        ready = pd.read_excel("_dane/dane.xlsx", index_col="Zlecenie")["Gotowe"] # Is order ready nie/tak
        searchData = self.panel_searchLineEdit.text() # Data from searching bar

        # Loading data to QListWidget
        for o in ordersIDs:
            if searchData in str(o):
                item = QtWidgets.QListWidgetItem(str(o))
                if ready[o] == "nie":
                    item.setBackground(QtGui.QColor("red"))
                elif ready[o] == "tak":
                    item.setBackground(QtGui.QColor("green"))
                self.panel_listWidget.addItem(item)

    def panel_openInfoWindow(self):
        try: self.currentOrderID = self.panel_listWidget.selectedItems()[0].text()
        except: pass
        
        #TODO: Sproawdx czy uzytkownik aktualnie nie jest na etapie dodawania lub edytowania zlecenia
        
        # Preparing window
        self.top_actualLabel.setText(f"Wybrane zlecenie: {self.currentOrderID}")
        self.stackedWidget.setCurrentIndex(1)

        # Load data from file to window
        df = pd.read_excel("_dane/dane.xlsx")
        data = df.loc[df["Zlecenie"] == int(self.currentOrderID)].reset_index()
        self.info_nrContentLabel.setText(str(data["Zlecenie"][0]))
        self.info_nrtelContentLabel.setText(str(data["Nrtel"][0]))
        self.info_modelContentLabel.setText(str(data["Model"][0]))
        self.info_dateContentLabel.setText(str(data["Data"][0]))
        self.info_descTextEdit.setText(str(data["Opis"][0]))
        self.info_usTextEdit.setText(str(data["Dlanas"][0]))

        if str(data["Gotowe"][0]) == "tak":
            self.info_isReadyCheckBox.setChecked(True)
        else:
            self.info_isReadyCheckBox.setChecked(False)
    
    def info_printButtonClicked(self):
        from docx import Document
        from docx.shared import Pt, Inches
        import os

        myfont =  Pt(14)

        document = Document()

        document.add_picture("mfcomp.png", width=Inches(3.5))
        pict = document.paragraphs[-1]
        pict.alignment = 1


        adres = document.add_paragraph("")
        adres_name = adres.add_run("37-100 Łańcut\nul. Rynek 31\ntel: 885 727 046")
        adres_name.italic = True
        adres.alignment = 1
        adres.paragraph_format.space_before = Inches(0.3)
        adres.paragraph_format.space_after = Inches(0.4)

        title = document.add_paragraph("")
        title_name = title.add_run("Potwierdzenie przyjęcia zlecenia numer: ")
        title_name.font.size = Pt(17)
        title_name.bold = True
        title_value = title.add_run(str(self.nr))
        title_value.font.size = Pt(17)
        title_value.bold = True
        title.alignment = 1
        #title_value.font.color = RGBColor(255,0,0)

        model = document.add_paragraph("")
        model_name = model.add_run("Model urządzenia: ")
        model_name.bold = True
        model_name.font.size = myfont
        model_value = model.add_run(str(self.model))
        model_value.font.size = myfont
        model.alignment = 1
        model.paragraph_format.space_before = Inches(1)


        data = document.add_paragraph("")
        data_title = data.add_run("Data przyjęcia: ")
        data_title.bold = True
        data_title.font.size = myfont
        data_value = data.add_run(str(self.date))
        data_value.font.size = myfont
        data.alignment = 1

        opis = document.add_paragraph("")
        opis_title = opis.add_run("Opis: ")
        opis_title.bold = True
        opis_title.font.size = myfont
        opis_value = opis.add_run(str(self.opis))
        opis_value.font.size = myfont
        opis.alignment = 1

        otherFont = Pt(16)

        data2 = document.add_paragraph("")
        data2_title = data2.add_run("Data odbioru: ")
        data2_title.bold = True
        data2_title.font.size = otherFont
        data2.paragraph_format.space_before = Inches(1)

        okres = document.add_paragraph("")
        okres_title = okres.add_run("Okres gwarancji: ")
        okres_title.bold = True
        okres_title.font.size = otherFont

        koszt = document.add_paragraph("")
        koszt_title = koszt.add_run("Koszt naprawy: ")
        koszt_title.bold = True
        koszt_title.font.size = otherFont

        document.save("druk.docx")

        os.startfile("druk.docx", "print")    

    def editButtonClicked(self):
        self.stackedWidget.setCurrentIndex(2)

        # Load data from file to window
        df = pd.read_excel("_dane/dane.xlsx")
        data = df.loc[df["Zlecenie"] == int(self.currentOrderID)].reset_index()
        self.edit_nrLineEdit.setText(str(data["Zlecenie"][0]))
        self.edit_nrtelLineEdit.setText(str(data["Nrtel"][0]))
        self.edit_modelLineEdit.setText(str(data["Model"][0]))
        self.edit_dateLineEdit.setText(str(data["Data"][0]))
        self.edit_descTextEdit.setText(str(data["Opis"][0]))
        self.edit_usTextEdit.setText(str(data["Dlanas"][0]))

        if str(data["Gotowe"][0]) == "tak":
            self.edit_isReadyCheckBox.setChecked(True)
        else:
            self.edit_isReadyCheckBox.setChecked(False)
    
    def info_removeButtonClicked(self):
        #TODO: Dokończ funkcjonalność przycisku
        pass

    def edit_saveButtonClicked(self):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Question)
        msg.setWindowTitle("Zapisywanie")
        msg.setText(f"Chcesz zapisać zlecenie nr {self.currentOrderID}?")
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        msg.setDefaultButton(QMessageBox.No)
        odp = msg.exec_()

        if odp == QtWidgets.QMessageBox.Yes:
            # Get data from app
            nrtel = self.edit_nrtelLineEdit.text()
            model = self.edit_modelLineEdit.text()
            desc = self.edit_descTextEdit.toPlainText()
            us = self.edit_usTextEdit.toPlainText()

            if self.edit_isReadyCheckBox.isChecked() == True:
                ready = "tak"
            else:
                ready = "nie"

            # Save changes to Excel file
            df = pd.read_excel("_dane/dane.xlsx")
            editIndex = df.index[df['Zlecenie'] == int(self.currentOrderID)].tolist()[0]
            df.loc[editIndex, ["Gotowe", "Nrtel", "Model", "Opis", "Dlanas"]] = [ready, nrtel, model, desc, us]
            df.to_excel("_dane/dane.xlsx", sheet_name="Sheet1", index=False)

            self.panel_openInfoWindow()
            self.loadListWidget()

    def edit_cancelButtonClicked(self):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Question)
        msg.setWindowTitle("Anulowanie")
        msg.setText(f"Chcesz anulować zapisywanie zlecenia nr. {self.currentOrderID}?")
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        msg.setDefaultButton(QMessageBox.No)
        odp = msg.exec_()

        if odp == QMessageBox.Yes:
            self.stackedWidget.setCurrentIndex(1)

    def panel_addButtonClicked(self):
        #TODO: jeżeli jest włączone okno edycji (2) to zapytaj czy chcesz zapisać edycje i przejść do dodawania

        def clearWindow():
            self.add_nrtelLineEdit.clear()
            self.add_modelLineEdit.clear()
            self.add_descTextEdit.clear()
            self.add_usTextEdit.clear()

        if self.stackedWidget.currentIndex() == 3: # Check if the user isn't already adding the order
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Question)
            msg.setWindowTitle("Czyszczenie")
            msg.setText(f"Chcesz wyczyścić aktualny postęp dodawania zlecenia?")
            msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
            msg.setDefaultButton(QMessageBox.No)
            odp = msg.exec_()
 
            if odp == QMessageBox.Yes:
                clearWindow()

        else: # If not open addPage
            self.top_actualLabel.setText(f"Dodawanie zlecenia nr. {self.currentOrderID}")
            self.stackedWidget.setCurrentIndex(3)
            clearWindow()

    def add_confirmClicked(self):
        from openpyxl import load_workbook
        from datetime import datetime

        nr = int(open("_dane/nr.txt", "r").readline())

        msg = QMessageBox()
        msg.setIcon(QMessageBox.Question)
        msg.setWindowTitle("Zapisywanie")
        msg.setText(f"Chcesz zapisać zlecenie nr {nr}?")
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        msg.setDefaultButton(QMessageBox.No)
        odp = msg.exec_()

        if odp == QtWidgets.QMessageBox.Yes:
            nrtel = self.add_nrtelLineEdit.text()
            model = self.add_modelLineEdit.text()
            desc = self.add_descTextEdit.toPlainText()
            us = self.add_usTextEdit.toPlainText()
            date = datetime.today().strftime("%d/%m/%Y")

            myFileName = "_dane/dane.xlsx"
            wb = load_workbook(filename=myFileName)
            ws = wb["Sheet1"]
            ws.append(["nie", nr, nrtel, model, date, desc, us])
            wb.save(filename=myFileName)
            wb.close()

            with open('_dane/nr.txt', "w") as file:
                file.write(str(nr+1))

            self.currentOrderID = nr
            self.stackedWidget.setCurrentIndex(0)
            self.loadListWidget()
            self.panel_openInfoWindow()

    def add_cancelClicked(self):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Question)
        msg.setWindowTitle("Anulowanie")
        msg.setText(f"Chcesz anulować zapisywanie zlecenia nr. {self.currentOrderID}?")
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        msg.setDefaultButton(QMessageBox.No)
        odp = msg.exec_()

        if odp == QMessageBox.Yes:
            self.top_actualLabel.setText(f"Strona główna")
            self.stackedWidget.setCurrentIndex(0)

    def panel_excelPreviewClicked(self):
        from excelPreview import Ui_ExcelPreview
        window = QDialog()
        ui = Ui_ExcelPreview()
        ui.setupUi(window)
        window.show()
        window.exec_()

